from flask import Flask, render_template_string, request, send_file
from docx import Document
import io
from datetime import datetime

app = Flask(__name__)

HTML_FORM = """
<!DOCTYPE html>
<html>
<head>
    <title>Redacción de demanda</title>
</head>
<body>
    <h1>Formulario para demanda ejecutiva de alimentos</h1>
    <form method='post'>
        <label>Nombre de la demandante:</label><br>
        <input type='text' name='demandante' required><br>

        <label>Nombre del NNA:</label><br>
        <input type='text' name='nna' required><br>

        <label>Nombre del demandado:</label><br>
        <input type='text' name='demandado' required><br>

        <label>Valor de la cuota alimentaria mensual ($):</label><br>
        <input type='number' name='cuota' required><br>

        <label>Fecha en que se dejó de pagar (YYYY-MM-DD):</label><br>
        <input type='date' name='fecha_ini' required><br><br>

        <input type='submit' value='Generar Demanda'>
    </form>
</body>
</html>
"""

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        # Obtener datos del formulario
        demandante = request.form['demandante']
        nna = request.form['nna']
        demandado = request.form['demandado']
        cuota = float(request.form['cuota'])
        fecha_ini_str = request.form.get('fecha_ini')

        # Validación de la fecha
        if not fecha_ini_str:
            return "⚠️ Error: Debes seleccionar una fecha de inicio para generar la demanda.", 400

        try:
            fecha_ini = datetime.strptime(fecha_ini_str, '%Y-%m-%d')
        except ValueError:
            return "⚠️ Error: El formato de fecha no es válido. Usa YYYY-MM-DD.", 400

        # Cálculo de intereses
        hoy = datetime.today()
        meses_adeudados = (hoy.year - fecha_ini.year) * 12 + hoy.month - fecha_ini.month
        interes = 0.005
        total_adeudado = sum([cuota * ((1 + interes) ** i) for i in range(meses_adeudados)])

        # Generar el documento Word
        doc = Document()
        doc.add_heading('Demanda Ejecutiva de Alimentos', 0)
        doc.add_paragraph(f'Demandante: {demandante}')
        doc.add_paragraph(f'Demandado: {demandado}')
        doc.add_paragraph(f'NNA beneficiario: {nna}')
        doc.add_paragraph(f'Fecha de inicio de mora: {fecha_ini.strftime("%Y-%m-%d")}')
        doc.add_paragraph(f'Cuota mensual: ${cuota:,.2f}')
        doc.add_paragraph(f'Total adeudado (con intereses): ${total_adeudado:,.2f}')

        # Preparar documento para descarga
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return send_file(file_stream, as_attachment=True, download_name='Demanda_Ejecutiva.docx')

    return render_template_string(HTML_FORM)

if __name__ == '__main__':
    app.run(debug=True)
